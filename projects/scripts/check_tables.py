#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word文档的表格结构
"""
from docx import Document
import sys

def check_tables(doc_path):
    """检查文档中的表格结构"""
    doc = Document(doc_path)
    print(f"\n检查文档: {doc_path}\n")
    print("=" * 80)

    print(f"总共有 {len(doc.tables)} 个表格\n")

    for table_idx, table in enumerate(doc.tables):
        print(f"表格 {table_idx + 1}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns) if table.rows else 0}")

        # 显示前5行的内容
        print(f"\n  前5行内容:")
        for row_idx, row in enumerate(table.rows[:5]):
            print(f"  行 {row_idx + 1}: ", end="")
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace('\n', ' ')[:30]
                print(f"[{cell_idx}:{cell_text}] ", end="")
            print()

        print()

    print("=" * 80)

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("用法: python3 check_tables.py <docx文件路径>")
        sys.exit(1)

    check_tables(sys.argv[1])
