#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Excel to Word Template Processor
读取 Excel 文件数据，填充到 Word 模板中
"""
import openpyxl
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import json
import os
from datetime import datetime

# 优先从环境变量读取配置，如果命令行参数已提供，则仍然沿用命令行参数。
# 这样可以在部署或 CI 环境中使用标准 Python os.getenv 方式配置变量。
EXCEL_ENV_PATH = os.getenv('EXCEL_PATH')
WORD_TEMPLATE_ENV_PATH = os.getenv('WORD_TEMPLATE_PATH')
OUTPUT_DIR_ENV = os.getenv('OUTPUT_DIR')


def format_currency(value):
    """格式化为货币格式（保留2位小数），前面添加￥符号"""
    if value is None:
        return "￥0.00"
    try:
        return f"￥{float(value):,.2f}"
    except (ValueError, TypeError):
        return "￥0.00"

def format_date(value):
    """格式化日期为年月日格式（去掉时间）"""
    if value is None:
        return ""
    try:
        if isinstance(value, datetime):
            return value.strftime("%Y年%m月%d日")
        elif isinstance(value, str):
            # 尝试解析日期字符串
            if " " in value:
                value = value.split(" ")[0]
            return value.replace("-", "年").replace("/", "年") + "日"
        return str(value)
    except:
        return str(value)

def set_run_bold(run, bold=True):
    """强制设置 run 的加粗格式"""
    run.font.bold = bold
    # 通过 XML 强制设置
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)

    # 创建或删除 b 元素
    b_element = rPr.find(qn('w:b'))
    if bold:
        if b_element is None:
            b_element = OxmlElement('w:b')
            rPr.append(b_element)
    else:
        if b_element is not None:
            rPr.remove(b_element)

def set_run_format(run, font_name='宋体', font_size=Pt(12), bold=False):
    """设置 run 的字体格式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    set_run_bold(run, bold=bold)

def set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False, alignment=None):
    """设置整个段落的字体格式，将所有 run 合并为一个"""
    # 收集所有文本
    full_text = paragraph.text

    # 清除所有 run
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # 创建一个新的 run
    new_run = paragraph.add_run(full_text)

    # 设置格式
    new_run.font.name = font_name
    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    new_run.font.size = font_size
    new_run.font.bold = bold

    # 通过 XML 强制设置加粗
    rPr = new_run._element.rPr
    if bold:
        b_element = rPr.find(qn('w:b'))
        if b_element is None:
            b_element = OxmlElement('w:b')
            rPr.append(b_element)
    else:
        b_element = rPr.find(qn('w:b'))
        if b_element is not None:
            rPr.remove(b_element)

    # 设置对齐
    if alignment is not None:
        paragraph.alignment = alignment

def process_excel_and_word(excel_path, word_template_path, output_dir):
    """
    处理 Excel 和 Word 文件

    Args:
        excel_path: Excel 文件路径
        word_template_path: Word 模板文件路径
        output_dir: 输出目录

    Returns:
        dict: 包含生成文件路径和数据的字典
    """
    # 读取 Excel 文件
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active

    # 提取关键单元格数据
    data = {
        'A1': ws['A1'].value or "",
        'A3': ws['A3'].value or "",
        'L3': format_date(ws['L3'].value),
        'C5': ws['C5'].value or "",
        'H6': ws['H6'].value or "",
        'J3': ws['J3'].value or "",
        'J6': ws['J6'].value or "",
        'R3': format_date(ws['R3'].value)
    }

    # 查找"未税合计"行，提取合计数据
    max_row = ws.max_row
    wsje_hj = None  # 未税合计
    se_hj = None    # 税额合计
    hsje_hj = None  # 含税合计

    for row_idx in range(10, max_row + 1):
        a_value = ws[f'A{row_idx}'].value

        # 找到"未税合计"行
        if a_value and '未税合计' in str(a_value):
            # 读取D、J、L列的数据
            wsje_hj = ws[f'D{row_idx}'].value  # D列（未税合计）
            se_hj = ws[f'J{row_idx}'].value    # J列（税额合计）
            hsje_hj = ws[f'L{row_idx}'].value  # L列（含税合计）
            break

    data['WSJEHJ'] = format_currency(wsje_hj)
    data['SEHJ'] = format_currency(se_hj)
    data['HSJEHJ'] = format_currency(hsje_hj)

    # 统计不同税率的数据（从第10行开始）
    sp_wsje = 0.0  # 13% 未税合计
    sp_se = 0.0    # 13% 税额合计
    sp_hsje = 0.0  # 13% 含税合计

    fw_wsje = 0.0  # 6% 未税合计
    fw_se = 0.0    # 6% 税额合计
    fw_hsje = 0.0  # 6% 含税合计

    for row_idx in range(10, max_row + 1):
        row = ws[row_idx]
        try:
            # 读取税率（第9列，I列，索引8）
            tax_rate = row[8].value  # I列（索引8）

            if tax_rate is None:
                continue

            # 解析税率
            if isinstance(tax_rate, str):
                if '%' in tax_rate:
                    tax_rate = float(tax_rate.replace('%', '')) / 100
                else:
                    tax_rate = float(tax_rate)
            else:
                tax_rate = float(tax_rate)

            # 读取金额
            wsje = float(row[9].value or 0)  # J列（索引9）金额不含税
            se = float(row[10].value or 0)    # K列（索引10）税额
            hsje = float(row[11].value or 0)  # L列（索引11）金额含税

            if abs(tax_rate - 0.13) < 0.01:
                sp_wsje += wsje
                sp_se += se
                sp_hsje += hsje
            elif abs(tax_rate - 0.06) < 0.01:
                fw_wsje += wsje
                fw_se += se
                fw_hsje += hsje

        except (ValueError, TypeError, IndexError):
            continue

    data['SPWSJE'] = format_currency(sp_wsje)
    data['SPSE'] = format_currency(sp_se)
    data['SPHSJE'] = format_currency(sp_hsje)

    data['FWWSJE'] = format_currency(fw_wsje)
    data['FWSE'] = format_currency(fw_se)
    data['FWHSJE'] = format_currency(fw_hsje)

    # 生成文件名：A1 + L3
    filename = f"{data['A1']}-{data['L3']}.docx"
    # 清理文件名中的非法字符
    filename = filename.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')

    # 读取 Word 模板
    doc = Document(word_template_path)

    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

    # 格式化文档：设置标题和报价单位名称的字体格式
    a1_count = 0  # 用于记录A1出现的次数
    total_a1_count = 0  # 统计A1总数
    a1_indices = []  # 记录A1出现的段落索引

    # 先统计A1出现的总次数和位置
    for idx, paragraph in enumerate(doc.paragraphs):
        if '{{A1}}' in paragraph.text or data['A1'] in paragraph.text:
            total_a1_count += 1
            a1_indices.append(idx)

    # 统计"报价一览表"和"报价明细表"的位置，用于判断页面
    ylb_index = None  # 报价一览表
    mxlb_index = None  # 报价明细表
    for idx, paragraph in enumerate(doc.paragraphs):
        if '报价一览表' in paragraph.text:
            ylb_index = idx
        if '报价明细表' in paragraph.text:
            mxlb_index = idx

    # 格式化各个段落
    for idx, paragraph in enumerate(doc.paragraphs):
        # 找到A1所在行（包含 {{A1}}）
        if '{{A1}}' in paragraph.text or data['A1'] in paragraph.text:
            a1_count += 1
            # 第一个A1设置为1号字加黑，第二个A1设置为小4号字不加黑
            if a1_count == 1:
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(26), bold=True)
            elif a1_count == 2:
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

        # 判断页面位置
        is_page1 = False  # 第1页
        is_page2 = False  # 第2页
        is_last_page = False  # 最后一页

        if ylb_index is not None and mxlb_index is not None:
            # 第1页：在"报价一览表"之前
            if idx < ylb_index:
                is_page1 = True
            # 第2页：在"报价一览表"之后，"报价明细表"之前
            elif ylb_index <= idx < mxlb_index:
                is_page2 = True
            # 最后一页：在"报价明细表"之后
            elif idx >= mxlb_index:
                is_last_page = True
        elif ylb_index is not None:
            # 只有"报价一览表"
            if idx < ylb_index:
                is_page1 = True
            else:
                is_last_page = True
        elif mxlb_index is not None:
            # 只有"报价明细表"
            if idx < mxlb_index:
                is_page1 = True
            else:
                is_last_page = True

        # 找到第一页L3所在行（包含 {{L3}} 或日期），排除第2页和最后一页
        if '{{L3}}' in paragraph.text or data['L3'] in paragraph.text or ('日' in paragraph.text and '月' in paragraph.text and '年' in paragraph.text):
            if is_page1:
                # 第一页L3：宋体4号字加黑
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(14), bold=True)
            else:
                # 第2页或最后一页L3：宋体小4号字不加粗
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

        # 找到报价单位名称所在行（包含 {{A3}} 或 "报价单位名称"）
        if '报价单位名称' in paragraph.text and ('{{A3}}' in paragraph.text or data['A3'] in paragraph.text):
            if is_page1:
                # 第一页A3：宋体4号字加黑并居中
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(14), bold=True,
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                # 第2页或最后一页A3：宋体小4号字不加粗
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

        # 找到其他包含 A3 数据的行（第2页和最后一页）
        elif data['A3'] in paragraph.text and '报价单位名称' not in paragraph.text:
            if is_page2 or is_last_page:
                # 第2页或最后一页A3：宋体小4号字不加粗
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

        # 找到C5所在行（包含 {{C5}}）
        if '{{C5}}' in paragraph.text or data['C5'] in paragraph.text:
            if is_page2 or is_last_page:
                # 第2页或最后一页C5：宋体小4号字不加粗
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

        # 找到H6所在行（包含 {{H6}}）
        if '{{H6}}' in paragraph.text or data['H6'] in paragraph.text:
            if is_page2 or is_last_page:
                # 第2页或最后一页H6：宋体小4号字不加粗
                set_paragraph_format(paragraph, font_name='宋体', font_size=Pt(12), bold=False)

    # 从Excel第10行开始读取数据并填充到Word表格
    # 读取Excel数据
    excel_data = []
    for row_idx in range(10, ws.max_row + 1):
        a_value = ws[f'A{row_idx}'].value

        # 检查是否到达"未税合计"行（检查A列）
        if a_value and '未税合计' in str(a_value):
            break

        # 读取当前行的数据（不读取BH字段）
        row_data = {
            'CPMC': ws[f'C{row_idx}'].value or '',
            'CPMS': ws[f'D{row_idx}'].value or '',
            'CPSL': ws[f'E{row_idx}'].value or '',
            'CPDW': ws[f'F{row_idx}'].value or '',
            'HSDJ': ws[f'G{row_idx}'].value or 0,
            'SL': ws[f'I{row_idx}'].value or '',
            'WSJE': ws[f'J{row_idx}'].value or 0,
            'HSJE': ws[f'L{row_idx}'].value or 0
        }
        excel_data.append(row_data)

    # 填充数据到Word表格
    # 使用第二个表格（表格2，从0开始索引为1）作为明细表
    if len(doc.tables) >= 2 and excel_data:
        detail_table = doc.tables[1]  # 第二个表格是明细表

        # 保留表头行（第一行），删除数据行
        # 从后往前删除，避免索引变化
        while len(detail_table.rows) > 1:
            # 删除最后一行
            row = detail_table.rows[-1]
            # 使用._element.remove()删除行
            row._element.getparent().remove(row._element)

        # 逐行添加Excel数据
        for bh_idx, row_data in enumerate(excel_data, start=1):
            # 添加新行
            new_row = detail_table.add_row()

            # 填充单元格数据
            # 列映射：Word列0=BH, Word列1=CPMC, Word列2=CPMS, Word列3=CPSL, Word列4=CPDW, Word列5=HSDJ, Word列6=SL, Word列7=WSJE, Word列8=HSJE
            cells = new_row.cells

            # 设置单元格内容
            cells[0].text = str(bh_idx)  # 编号（使用循环变量）
            cells[1].text = str(row_data['CPMC'])  # 产品名称
            cells[2].text = str(row_data['CPMS'])  # 产品描述
            cells[3].text = str(row_data['CPSL'])  # 数量
            cells[4].text = str(row_data['CPDW'])  # 单位

            # 含税单价（添加￥符号）
            try:
                hsdj_value = float(row_data['HSDJ'])
                cells[5].text = f"￥{hsdj_value:,.2f}"
            except (ValueError, TypeError):
                cells[5].text = str(row_data['HSDJ'])

            # 税率
            sl_value = row_data['SL']
            if isinstance(sl_value, (int, float)):
                cells[6].text = f"{sl_value*100:.0f}%" if sl_value < 1 else f"{sl_value}%"
            else:
                cells[6].text = str(sl_value)

            # 未税金额（添加￥符号）
            try:
                wsje_value = float(row_data['WSJE'])
                cells[7].text = f"￥{wsje_value:,.2f}"
            except (ValueError, TypeError):
                cells[7].text = str(row_data['WSJE'])

            # 含税金额（添加￥符号）
            try:
                hsje_value = float(row_data['HSJE'])
                cells[8].text = f"￥{hsje_value:,.2f}"
            except (ValueError, TypeError):
                cells[8].text = str(row_data['HSJE'])

            # 设置新行所有单元格的字体为小五号宋体（9磅）
            for cell in new_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(9)  # 小五号字

    # 格式化第一页的表格：统一列宽、调整字体、设置边框样式
    # 只处理第一个表格，确保只影响第一页
    if doc.tables:  # 检查是否有表格
        first_table = doc.tables[0]  # 只处理第一个表格（第一页的表格）
        for row in first_table.rows:
            for cell in row.cells:
                # 设置字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)

    # 保存文件
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return {
        'success': True,
        'output_file': filename,
        'output_path': output_path,
        'data': data
    }

def _resolve_paths_from_args_or_env():
    """优先使用命令行参数，缺失时回退到环境变量。"""
    excel_path = sys.argv[1] if len(sys.argv) > 1 else EXCEL_ENV_PATH
    word_template_path = sys.argv[2] if len(sys.argv) > 2 else WORD_TEMPLATE_ENV_PATH
    output_dir = sys.argv[3] if len(sys.argv) > 3 else OUTPUT_DIR_ENV

    return excel_path, word_template_path, output_dir


if __name__ == '__main__':
    excel_path, word_template_path, output_dir = _resolve_paths_from_args_or_env()

    if not excel_path or not word_template_path or not output_dir:
        print(json.dumps({'error': 'Missing arguments or environment variables. Set EXCEL_PATH, WORD_TEMPLATE_PATH and OUTPUT_DIR, or pass them as command-line arguments.'}))
        sys.exit(1)

    try:
        result = process_excel_and_word(excel_path, word_template_path, output_dir)
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)
