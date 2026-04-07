#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word表格的字体大小
"""
from docx import Document
from docx.oxml.ns import qn
import sys

def check_table_font_size(doc_path):
    """检查Word表格的字体大小"""
    doc = Document(doc_path)
    print(f"\n检查文档: {doc_path}\n")
    print("=" * 80)

    # 检查第二个表格（明细表）
    if len(doc.tables) >= 2:
        detail_table = doc.tables[1]
        print(f"表格2（明细表）字体大小：\n")

        for row_idx, row in enumerate(detail_table.rows):
            if row_idx == 0:
                print(f"表头行（行{row_idx+1}）：")
            else:
                print(f"数据行{row_idx}（行{row_idx+1}）：")

            for col_idx, cell in enumerate(row.cells):
                # 检查第一个段落的字体大小
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if para.runs:
                        run = para.runs[0]
                        font_name = run.font.name
                        font_size = run.font.size
                        if font_size:
                            size_pt = font_size.pt
                        else:
                            size_pt = None
                        print(f"  列{col_idx}: {font_name}, {size_pt}磅")

            print()

    print("=" * 80)

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("用法: python3 check_table_font.py <docx文件路径>")
        sys.exit(1)

    check_table_font_size(sys.argv[1])
